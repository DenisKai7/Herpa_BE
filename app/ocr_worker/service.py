"""OCR worker extraction implementation.

This module intentionally owns heavy imports: torch, transformers, PyMuPDF, Pillow,
pandas, and document parsers. The main backend must not import this module.
"""

from __future__ import annotations

import asyncio
import io
import os
import re
import time
from typing import Any

from pydantic import BaseModel, Field

from app.core.config import settings


PDF_EXTENSIONS = {"pdf"}
IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "bmp", "tiff"}
TEXT_EXTENSIONS = {"txt", "md", "json"}
DOCX_EXTENSIONS = {"docx"}
EXCEL_CSV_EXTENSIONS = {"xlsx", "xls", "csv"}
ALL_ALLOWED_EXTENSIONS = PDF_EXTENSIONS | IMAGE_EXTENSIONS | TEXT_EXTENSIONS | DOCX_EXTENSIONS | EXCEL_CSV_EXTENSIONS

MIME_BY_EXT = {
    "pdf": "application/pdf",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "tiff": "image/tiff",
    "txt": "text/plain",
    "md": "text/plain",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "xls": "application/vnd.ms-excel",
    "csv": "text/csv",
}


class WorkerOcrResult(BaseModel):
    success: bool
    extracted_text: str = ""
    raw_text: str = ""
    normalized_text: str = ""
    detected_type: str = "unknown"
    visible_labels: list[str] = Field(default_factory=list)
    chemical_terms: list[str] = Field(default_factory=list)
    molecular_formulas: list[str] = Field(default_factory=list)
    numeric_labels: list[str] = Field(default_factory=list)
    document_sections: list[str] = Field(default_factory=list)
    tables: list[str] = Field(default_factory=list)
    confidence: float = 0.0
    warnings: list[str] = Field(default_factory=list)
    model_id: str = "ocr-worker"
    method: str = "parser"
    processing_ms: int = 0


def check_runtime() -> dict[str, Any]:
    cache_dir = os.environ.get("HF_HOME", "/models/huggingface")
    os.makedirs(cache_dir, exist_ok=True)
    imports: dict[str, bool] = {}
    for module_name in ("fitz", "PIL", "torch", "transformers"):
        try:
            __import__(module_name)
            imports[module_name] = True
        except Exception:
            imports[module_name] = False
    return {
        "status": "ok",
        "service": "medical_ocr_worker",
        "cache_dir": cache_dir,
        "cache_dir_exists": os.path.isdir(cache_dir),
        "imports": imports,
        "lazy_load": settings.OCR_LAZY_LOAD,
    }


def _ext(filename: str) -> str:
    if "." not in filename:
        raise ValueError("missing_file_extension")
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ALL_ALLOWED_EXTENSIONS:
        raise ValueError(f"unsupported_file_extension:{ext}")
    return ext


def _sniff_mime(content: bytes) -> str:
    head = content[:4096]
    if head.startswith(b"%PDF"):
        return "application/pdf"
    if head.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head.startswith(b"RIFF") and b"WEBP" in head[:16]:
        return "image/webp"
    if head.startswith(b"BM"):
        return "image/bmp"
    if head.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if content[:4] == b"PK\x03\x04":
        import zipfile

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = set(zf.namelist())
                if "word/document.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if "xl/workbook.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except Exception:
            pass
    try:
        content[:8192].decode("utf-8")
        return "text/plain"
    except UnicodeDecodeError:
        return "application/octet-stream"


def _validate(filename: str, mime_type: str, content: bytes) -> str:
    ext = _ext(filename)
    sniffed = _sniff_mime(content)
    expected = MIME_BY_EXT[ext]
    compatible = sniffed == expected or (ext == "md" and sniffed == "text/plain") or (ext == "csv" and sniffed == "text/plain")
    if not compatible:
        raise ValueError(f"mime_extension_mismatch:{sniffed}:{ext}")
    return ext


def _classify(text: str) -> dict[str, Any]:
    chemical_symbols = list(set(re.findall(r"\b(OH|COOH|NH2|CH3|OCH3|CH2|C\d+H\d+|H2O|CO2)\b", text)))
    molecular_formulas = list(set(re.findall(r"\b[A-Z][A-Za-z]{0,2}\d+[A-Za-z0-9]*\b", text)))
    chemical_terms = list(set(re.findall(r"\b(flavonoid|alkaloid|saponin|tannin|curcumin|gingerol|xanthorrhizol|terpenoid|steroid|glycoside|phenolic)\b", text.lower())))
    numeric_labels = list(set(re.findall(r"\b\d{1,2}\b", text)))
    has_table = "|" in text and text.count("|") > 4
    tables = [line for line in text.splitlines() if "|" in line] if has_table else []
    sections = [line.strip() for line in text.splitlines() if re.match(r"^(bab|chapter|section|abstrak|abstract|daftar|tabel|gambar|\d+\.)", line.lower().strip())]
    detected_type = "plain_text" if text.strip() else "unknown"
    if has_table:
        detected_type = "table"
    elif len(chemical_symbols) > 1 or molecular_formulas:
        detected_type = "chemical_structure_diagram"
    elif len(text.split()) > 100:
        detected_type = "scanned_document"
    return {
        "detected_type": detected_type,
        "visible_labels": chemical_symbols,
        "chemical_terms": chemical_terms,
        "molecular_formulas": molecular_formulas,
        "numeric_labels": numeric_labels,
        "document_sections": sections,
        "tables": tables,
    }


def _normalize(text: str) -> str:
    cleaned = " ".join(text.split())
    return cleaned[: settings.ATTACHMENT_CONTEXT_MAX_CHARS]


def _text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _pdf_text(content: bytes) -> str:
    import fitz

    document = fitz.open(stream=content, filetype="pdf")
    try:
        chunks = []
        for index in range(min(len(document), settings.ATTACHMENT_MAX_PDF_PAGES)):
            page = document[index]
            text = page.get_text("text").strip()
            if text:
                chunks.append(f"[Page {index + 1}]:\n{text}")
        return "\n\n".join(chunks)
    finally:
        document.close()


def _docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                chunks.append(" | ".join(row_text))
    return "\n\n".join(chunks)


def _excel_csv(content: bytes, filename: str) -> str:
    import pandas as pd

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "csv":
        df = pd.read_csv(io.BytesIO(content))
    else:
        df = pd.read_excel(io.BytesIO(content))
    try:
        return df.to_markdown(index=False)
    except Exception:
        return df.to_string(index=False)


def _image(content: bytes):
    from PIL import Image, ImageOps

    Image.MAX_IMAGE_PIXELS = settings.OCR_MAX_IMAGE_PIXELS
    with Image.open(io.BytesIO(content)) as verifier:
        verifier.verify()
    image = Image.open(io.BytesIO(content))
    image = ImageOps.exif_transpose(image)
    image = image.convert("RGB")
    if image.width * image.height > settings.OCR_MAX_IMAGE_PIXELS:
        image.thumbnail((4096, 4096))
    return image


def _render_pdf_pages(content: bytes) -> list[Any]:
    import fitz
    from PIL import Image

    document = fitz.open(stream=content, filetype="pdf")
    images = []
    try:
        for index in range(min(len(document), settings.ATTACHMENT_OCR_PAGE_LIMIT)):
            page = document[index]
            pix = page.get_pixmap(dpi=150)
            images.append(Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB"))
    finally:
        document.close()
    return images


def _resolve_device() -> str:
    import torch

    configured = settings.OCR_DEVICE.lower()
    if configured != "auto":
        return configured
    return "cuda" if torch.cuda.is_available() else "cpu"


def _resolve_dtype(device: str):
    import torch

    if settings.OCR_DTYPE == "float16":
        return torch.float16
    if settings.OCR_DTYPE == "bfloat16":
        return torch.bfloat16
    return torch.float32


class GotOcrService:
    def __init__(self) -> None:
        self._processor = None
        self._model = None
        self._device = None
        self._lock = asyncio.Lock()

    async def ensure_loaded(self) -> None:
        if self._processor is not None and self._model is not None:
            return
        async with self._lock:
            if self._processor is not None and self._model is not None:
                return
            from transformers import AutoProcessor, GotOcr2ForConditionalGeneration

            import torch

            device = _resolve_device()
            dtype = _resolve_dtype(device)
            processor = AutoProcessor.from_pretrained(settings.OCR_MODEL_ID, token=settings.HF_API_TOKEN)
            model = GotOcr2ForConditionalGeneration.from_pretrained(
                settings.OCR_MODEL_ID,
                token=settings.HF_API_TOKEN,
                torch_dtype=dtype,
                low_cpu_mem_usage=True,
            )
            model.to(device)
            model.eval()
            self._processor = processor
            self._model = model
            self._device = device

    async def extract_image(self, image) -> str:
        import torch

        await self.ensure_loaded()
        loop = asyncio.get_running_loop()

        def _run() -> str:
            with torch.inference_mode():
                inputs = self._processor(images=image, return_tensors="pt")
                inputs = {key: value.to(self._device) for key, value in inputs.items()}
                generated = self._model.generate(**inputs, max_new_tokens=settings.OCR_MAX_NEW_TOKENS)
                decoded = self._processor.batch_decode(generated, skip_special_tokens=True)
                return decoded[0] if decoded else ""

        return await asyncio.wait_for(loop.run_in_executor(None, _run), timeout=settings.OCR_TIMEOUT_SECONDS)


got_ocr_service = GotOcrService()


async def extract_file(*, filename: str, mime_type: str, content: bytes) -> dict[str, Any]:
    start = time.time()
    ext = _validate(filename, mime_type, content)
    method = "parser"
    warnings: list[str] = []

    if ext in TEXT_EXTENSIONS:
        raw_text = _text(content)
    elif ext in PDF_EXTENSIONS:
        raw_text = _pdf_text(content)
        if len(raw_text.strip()) < 40:
            method = "GOT-OCR2"
            rendered = _render_pdf_pages(content)
            ocr_chunks = []
            for image in rendered:
                extracted = await got_ocr_service.extract_image(image)
                if extracted.strip():
                    ocr_chunks.append(extracted)
            raw_text = "\n\n".join(ocr_chunks)
    elif ext in IMAGE_EXTENSIONS:
        method = "GOT-OCR2"
        image = _image(content)
        raw_text = await got_ocr_service.extract_image(image)
    elif ext in DOCX_EXTENSIONS:
        raw_text = _docx(content)
    elif ext in EXCEL_CSV_EXTENSIONS:
        raw_text = _excel_csv(content, filename)
        method = "table-parser"
    else:
        raise ValueError(f"unsupported_file_extension:{ext}")

    normalized = _normalize(raw_text)
    classified = _classify(raw_text)
    confidence = 0.9 if normalized else 0.0
    result = WorkerOcrResult(
        success=bool(normalized),
        extracted_text=normalized,
        raw_text=normalized,
        normalized_text=normalized,
        detected_type=classified["detected_type"],
        visible_labels=classified["visible_labels"],
        chemical_terms=classified["chemical_terms"],
        molecular_formulas=classified["molecular_formulas"],
        numeric_labels=classified["numeric_labels"],
        document_sections=classified["document_sections"],
        tables=classified["tables"],
        confidence=confidence,
        warnings=warnings,
        model_id=settings.OCR_MODEL_ID if method == "GOT-OCR2" else "parser",
        method=method,
        processing_ms=int((time.time() - start) * 1000),
    )
    return result.model_dump()
