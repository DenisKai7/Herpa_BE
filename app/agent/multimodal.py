from __future__ import annotations

import base64
import hashlib
import io
import json
import logging
import re
import time
from typing import Any, Final, Optional

from pydantic import BaseModel, Field, ValidationError

from app.core.config import settings
from app.core.huggingface_vlm_client import (
    HuggingFaceVlmClient,
    HuggingFaceVlmError,
    VlmModelsUnavailableError,
    VlmModelRoute,
    availability_cache,
    resolve_vlm_candidates,
    preprocess_image as client_preprocess_image,
)

logger = logging.getLogger(__name__)

_vlm_client: HuggingFaceVlmClient | None = None


def set_vlm_client(client: HuggingFaceVlmClient | None) -> None:
    global _vlm_client
    _vlm_client = client


def get_vlm_client() -> HuggingFaceVlmClient:
    global _vlm_client
    if _vlm_client is None:
        _vlm_client = HuggingFaceVlmClient(settings)
    return _vlm_client


PDF_EXTENSIONS: Final[frozenset[str]] = frozenset({"pdf"})
IMAGE_EXTENSIONS: Final[frozenset[str]] = frozenset({"jpg", "jpeg", "png", "webp", "bmp", "tiff"})
TEXT_EXTENSIONS: Final[frozenset[str]] = frozenset({"txt", "md", "json"})
DOCX_EXTENSIONS: Final[frozenset[str]] = frozenset({"docx"})
EXCEL_CSV_EXTENSIONS: Final[frozenset[str]] = frozenset({"xlsx", "xls", "csv"})
ALL_ALLOWED_EXTENSIONS: Final[frozenset[str]] = PDF_EXTENSIONS | IMAGE_EXTENSIONS | TEXT_EXTENSIONS | DOCX_EXTENSIONS | EXCEL_CSV_EXTENSIONS
ALLOWED_MIME_TYPES: Final[dict[str, frozenset[str]]] = {
    "application/pdf": PDF_EXTENSIONS,
    "image/jpeg": frozenset({"jpg", "jpeg"}),
    "image/png": frozenset({"png"}),
    "image/webp": frozenset({"webp"}),
    "image/bmp": frozenset({"bmp"}),
    "image/tiff": frozenset({"tiff"}),
    "text/plain": frozenset({"txt", "md"}),
    "application/json": frozenset({"json"}),
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCX_EXTENSIONS,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": frozenset({"xlsx"}),
    "application/vnd.ms-excel": frozenset({"xls"}),
    "text/csv": frozenset({"csv"}),
}


class VisualExtractionResult(BaseModel):
    success: bool
    raw_text: str = ""
    normalized_text: str = ""
    detected_type: str = "unknown"
    visible_labels: list[str] = Field(default_factory=list)
    chemical_terms: list[str] = Field(default_factory=list)
    molecular_formulas: list[str] = Field(default_factory=list)
    plant_names: list[str] = Field(default_factory=list)
    claims: list[str] = Field(default_factory=list)
    uncertainties: list[str] = Field(default_factory=list)
    numeric_labels: list[str] = Field(default_factory=list)
    document_sections: list[str] = Field(default_factory=list)
    tables: list[str] = Field(default_factory=list)
    confidence: float = 0.0
    warnings: list[str] = Field(default_factory=list)
    model_id: str = settings.VLM_PRIMARY_MODEL
    processing_ms: int = 0
    requested_model: Optional[str] = None
    used_model: Optional[str] = None
    fallback_used: bool = False
    fallback_reason: Optional[str] = None


OcrExtractionResult = VisualExtractionResult


class VlmStructuredOutput(BaseModel):
    detected_content_type: str = "unknown"
    visual_description: str = ""
    extracted_text: str = ""
    plant_names: list[str] = Field(default_factory=list)
    compound_names: list[str] = Field(default_factory=list)
    molecular_formulas: list[str] = Field(default_factory=list)
    visible_labels: list[str] = Field(default_factory=list)
    claims: list[str] = Field(default_factory=list)
    uncertainties: list[str] = Field(default_factory=list)
    confidence: float = 0.0


class AttachmentAnalysisResult(BaseModel):
    attachment_id: Optional[str] = None
    filename: str
    mime_type: str
    file_sha256: str
    extraction_method: str
    extracted_text: str
    structured_data: dict = Field(default_factory=dict)
    detected_entities: list[dict] = Field(default_factory=list)
    neo4j_candidates: list[dict] = Field(default_factory=list)
    verification_status: str = "unverified"
    confidence: float = 0.0
    warnings: list[str] = Field(default_factory=list)
    processing_ms: int


def validate_file_type(filename: str, content_type: str | None = None, content: bytes | None = None) -> str:
    if "." not in filename:
        raise ValueError(f"File '{filename}' tidak memiliki ekstensi.")
    ext = filename.rsplit(".", maxsplit=1)[-1].lower()
    if ext not in ALL_ALLOWED_EXTENSIONS:
        raise ValueError(f"Format file '.{ext}' tidak didukung.")
    if content is not None:
        sniffed_mime = sniff_mime_type(content, filename)
        if sniffed_mime not in ALLOWED_MIME_TYPES:
            raise ValueError(f"MIME type dari isi file '{sniffed_mime}' tidak diizinkan.")
        if ext not in ALLOWED_MIME_TYPES[sniffed_mime]:
            raise ValueError(f"MIME type dari isi file '{sniffed_mime}' tidak cocok dengan ekstensi '.{ext}'.")
    if content_type:
        normalized_mime = content_type.split(";")[0].strip().lower()
        if normalized_mime not in ALLOWED_MIME_TYPES:
            raise ValueError(f"MIME type '{normalized_mime}' tidak diizinkan.")
        if ext not in ALLOWED_MIME_TYPES[normalized_mime]:
            raise ValueError(f"MIME type '{normalized_mime}' tidak cocok dengan ekstensi '.{ext}'.")
    return ext


def sniff_mime_type(content: bytes, filename: str = "") -> str:
    head = content[:4096]
    if head.startswith(b"%PDF"):
        return "application/pdf"
    if head.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head.startswith(b"RIFF") and b"WEBP" in head[:16]:
        return "image/webp"
    if head.startswith(b"BM"):
        return "image/bmp"
    if head.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if content[:4] == b"PK\x03\x04":
        import zipfile
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = set(zf.namelist())
                if "word/document.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if "xl/workbook.xml" in names:
                    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        except Exception:
            pass
    try:
        content[:8192].decode("utf-8")
        return "text/plain"
    except UnicodeDecodeError:
        return "application/octet-stream"


def preprocess_image(data: bytes) -> tuple[bytes, str]:
    try:
        return client_preprocess_image(data, max_pixels=settings.VLM_MAX_IMAGE_PIXELS)
    except Exception as exc:
        raise ValueError("corrupt_image") from exc


def to_data_uri(content: bytes, mime_type: str) -> str:
    encoded = base64.b64encode(content).decode("ascii")
    if len(encoded) > settings.VLM_MAX_BASE64_BYTES:
        raise ValueError("vlm_payload_too_large")
    return f"data:{mime_type};base64,{encoded}"


def calculate_sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _extract_from_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _strip_markdown_json(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?\s*", "", stripped, flags=re.IGNORECASE)
        stripped = re.sub(r"\s*```$", "", stripped)
    match = re.search(r"\{.*\}", stripped, flags=re.DOTALL)
    return match.group(0) if match else stripped


def parse_vlm_json(text: str) -> VlmStructuredOutput:
    data = json.loads(_strip_markdown_json(text))
    return VlmStructuredOutput.model_validate(data)


def classify_extracted_text(text: str) -> dict[str, Any]:
    word_count = len(text.split())
    chemical_symbols = list(set(re.findall(r"\b(OH|COOH|NH2|CH3|OCH3|CH2|C\d+H\d+|H2O|CO2)\b", text)))
    molecular_formulas = list(set(re.findall(r"\b[CNO SPHFIBrcl]{1,8}\d+[A-Za-z0-9]*\b", text)))
    molecular_formulas = [f for f in molecular_formulas if any(c.isdigit() for c in f) and len(f) > 2]
    chemical_terms = list(set(re.findall(r"\b(flavonoid|alkaloid|saponin|tannin|curcumin|gingerol|xanthorrhizol|terpenoid|steroid|glycoside|phenolic)\b", text.lower())))
    numeric_labels = list(set(re.findall(r"\b\d{1,2}\b", text)))
    has_table = "|" in text and text.count("|") > 4
    tables = [line for line in text.splitlines() if "|" in line] if has_table else []
    sections = [line.strip() for line in text.splitlines() if re.match(r"^(bab|chapter|section|abstrak|abstract|daftar|tabel|gambar|\d+\.)", line.lower().strip())]
    detected_type = "unknown"
    if has_table or any(k in text.lower() for k in ["table", "tabel"]):
        detected_type = "table"
    elif len(chemical_symbols) > 1 or molecular_formulas or any(k in text.lower() for k in ["benzene", "cincin", "skeletal", "structure"]):
        detected_type = "chemical_structure_diagram"
    elif word_count > 100:
        detected_type = "scanned_document"
    elif word_count > 0:
        detected_type = "plain_text"
    return {"detected_type": detected_type, "chemical_symbols": chemical_symbols, "molecular_formulas": molecular_formulas, "chemical_terms": chemical_terms, "numeric_labels": numeric_labels, "document_sections": sections, "tables": tables}


def _analysis_from_visual_result(*, filename: str, mime_type: str, content: bytes, result: VisualExtractionResult, method: str) -> AttachmentAnalysisResult:
    heuristics = classify_extracted_text(result.raw_text)
    structured_data = {
        "visible_labels": result.visible_labels or heuristics["chemical_symbols"],
        "chemical_terms": result.chemical_terms or heuristics["chemical_terms"],
        "compound_names": result.chemical_terms or heuristics["chemical_terms"],
        "plant_names": result.plant_names,
        "molecular_formulas": result.molecular_formulas or heuristics["molecular_formulas"],
        "numeric_labels": result.numeric_labels or heuristics["numeric_labels"],
        "document_sections": result.document_sections or heuristics["document_sections"],
        "tables": result.tables or heuristics["tables"],
        "claims": result.claims,
        "uncertainties": result.uncertainties,
        "detected_type": result.detected_type if result.detected_type != "unknown" else heuristics["detected_type"],
        "requested_model": result.requested_model,
        "used_model": result.used_model,
        "fallback_used": result.fallback_used,
        "fallback_reason": result.fallback_reason,
    }
    return AttachmentAnalysisResult(
        filename=filename,
        mime_type=mime_type,
        file_sha256=calculate_sha256(content),
        extraction_method=method,
        extracted_text=result.raw_text[: settings.ATTACHMENT_CONTEXT_MAX_CHARS],
        structured_data=structured_data,
        detected_entities=[{"type": "plant_name", "value": value} for value in structured_data["plant_names"]]
        + [{"type": "compound_name", "value": value} for value in structured_data["compound_names"]]
        + [{"type": "molecular_formula", "value": value} for value in structured_data["molecular_formulas"]],
        verification_status="not_applicable",
        confidence=result.confidence,
        warnings=result.warnings,
        processing_ms=result.processing_ms,
    )


def _vlm_system_prompt() -> str:
    return (
        "Anda adalah analis visual medis/herbal/farmasi. Ekstrak hanya bukti yang terlihat pada gambar. "
        "Jangan mengarang identitas tanaman, senyawa, formula, diagnosis, atau klaim klinis. "
        "Balas JSON valid saja dengan field: detected_content_type, visual_description, extracted_text, "
        "plant_names, compound_names, molecular_formulas, visible_labels, claims, uncertainties, confidence."
    )


async def analyze_with_fallback(
    *,
    image_bytes: bytes,
    mime_type: str,
    user_question: str,
    system_prompt: str,
    requested_model: str | None = None,
) -> dict:
    vlm_client = get_vlm_client()
    route = resolve_vlm_candidates(requested_model)
    failures: list[dict] = []
    start_time = time.time()

    for model_id in route.candidate_models:
        if availability_cache.is_unavailable(model_id):
            continue

        try:
            result = await vlm_client.analyze_image(
                model_id=model_id,
                image_bytes=image_bytes,
                mime_type=mime_type,
                question=user_question,
                system_prompt=system_prompt,
            )

            latency_ms = int((time.time() - start_time) * 1000)
            fallback_used = model_id != (requested_model or route.candidate_models[0])
            logger.info(
                "vlm_request_completed requested_model=%s used_model=%s fallback_used=%s latency_ms=%d",
                requested_model or route.candidate_models[0],
                model_id,
                str(fallback_used).lower(),
                latency_ms,
            )

            return {
                **result,
                "requested_model": requested_model or route.candidate_models[0],
                "used_model": model_id,
                "fallback_used": fallback_used,
                "model_failures": failures,
            }

        except HuggingFaceVlmError as exc:
            failures.append({
                "model_id": model_id,
                "code": exc.code,
            })

            if exc.code == "model_not_supported":
                availability_cache.mark_unavailable(
                    model_id=model_id,
                    reason=exc.code,
                    ttl=settings.VLM_FAILURE_COOLDOWN_SECONDS,
                )
                logger.warning(
                    "vlm_model_unavailable model_id=%s reason=%s cooldown_seconds=%d",
                    model_id,
                    exc.code,
                    settings.VLM_FAILURE_COOLDOWN_SECONDS,
                )
            else:
                if not exc.retryable:
                    availability_cache.mark_unavailable(
                        model_id=model_id,
                        reason=exc.code,
                        ttl=settings.VLM_FAILURE_COOLDOWN_SECONDS,
                    )
                logger.warning(
                    "Remote VLM error: model=%s code=%s trying_next_candidate=true",
                    model_id,
                    exc.code,
                )

            continue
        except Exception as exc:
            logger.exception(
                "Unexpected error on remote VLM model=%s",
                model_id,
            )
            failures.append({
                "model_id": model_id,
                "code": "unexpected_error",
            })
            availability_cache.mark_unavailable(
                model_id=model_id,
                reason="unexpected_error",
                ttl=settings.VLM_FAILURE_COOLDOWN_SECONDS,
            )
            continue

    raise VlmModelsUnavailableError(
        "Seluruh remote VLM tidak tersedia.",
        failures=failures,
    )


async def _analyze_images_with_vlm(*, filename: str, mime_type: str, content: bytes, image_bytes: bytes, image_mime: str, user_query: str | None) -> AttachmentAnalysisResult:
    client = get_vlm_client()
    start = time.time()
    try:
        question = user_query or "Analisis gambar ini sesuai bukti visual. Keluarkan JSON valid sesuai schema."
        response = await analyze_with_fallback(
            image_bytes=image_bytes,
            mime_type=image_mime,
            user_question=question,
            system_prompt=_vlm_system_prompt(),
        )
        raw = str(response.get("content") or "")
        try:
            structured = parse_vlm_json(raw)
        except (json.JSONDecodeError, ValidationError):
            repaired = await client.repair_json(
                model_id=response.get("used_model"),
                raw_text=raw,
                system_prompt="Anda memperbaiki output menjadi JSON valid sesuai schema visual."
            )
            structured = parse_vlm_json(repaired)

        fallback_reason = None
        if response.get("fallback_used"):
            failures = response.get("model_failures") or []
            if failures:
                fallback_reason = failures[0].get("code")

        visual = VisualExtractionResult(
            success=True,
            raw_text="\n".join(part for part in [structured.visual_description, structured.extracted_text] if part),
            normalized_text=structured.extracted_text,
            detected_type=structured.detected_content_type or "unknown",
            visible_labels=structured.visible_labels,
            chemical_terms=structured.compound_names,
            molecular_formulas=structured.molecular_formulas,
            plant_names=structured.plant_names,
            claims=structured.claims,
            uncertainties=structured.uncertainties,
            confidence=max(0.0, min(1.0, structured.confidence)),
            model_id=str(response.get("used_model") or settings.VLM_PRIMARY_MODEL),
            processing_ms=int((time.time() - start) * 1000),
            requested_model=response.get("requested_model"),
            used_model=response.get("used_model"),
            fallback_used=bool(response.get("fallback_used")),
            fallback_reason=fallback_reason,
        )
        return _analysis_from_visual_result(filename=filename, mime_type=mime_type, content=content, result=visual, method="hf-vlm")
    except (HuggingFaceVlmError, VlmModelsUnavailableError) as exc:
        if isinstance(exc, HuggingFaceVlmError):
            logger.error("HuggingFaceVlmError during VLM image analysis: code=%s message=%s", exc.code, exc.message)
        else:
            logger.error("VlmModelsUnavailableError during VLM image analysis: message=%s", exc.message)
        raise exc


def _extract_docx_text(content: bytes) -> str:
    from docx import Document
    document = Document(io.BytesIO(content))
    lines = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _extract_xlsx_csv_text(content: bytes, ext: str) -> str:
    if ext == "csv":
        return _extract_from_text(content)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    rows: list[str] = []
    for ws in wb.worksheets[:3]:
        rows.append(f"Sheet: {ws.title}")
        for row in ws.iter_rows(max_row=50, max_col=20, values_only=True):
            rows.append(" | ".join("" if v is None else str(v) for v in row))
    return "\n".join(rows)


def _extract_pdf_text(content: bytes) -> str:
    import fitz
    doc = fitz.open(stream=content, filetype="pdf")
    text_parts = []
    for index in range(min(len(doc), settings.ATTACHMENT_MAX_PDF_PAGES)):
        text_parts.append(doc[index].get_text("text"))
    return "\n".join(text_parts)


async def process_attachment(*, filename: str, mime_type: str, content: bytes, user_query: str | None) -> AttachmentAnalysisResult:
    start_time = time.time()
    ext = validate_file_type(filename, mime_type, content)
    if len(content) > settings.VLM_MAX_FILE_SIZE_MB * 1024 * 1024 and ext in IMAGE_EXTENSIONS:
        raise ValueError("vlm_file_too_large")

    if ext in TEXT_EXTENSIONS:
        text = _extract_from_text(content)
    elif ext in DOCX_EXTENSIONS:
        text = _extract_docx_text(content)
    elif ext in EXCEL_CSV_EXTENSIONS:
        text = _extract_xlsx_csv_text(content, ext)
    elif ext in PDF_EXTENSIONS:
        text = _extract_pdf_text(content)
    else:
        text = ""

    if ext in IMAGE_EXTENSIONS:
        safe_bytes, safe_mime = preprocess_image(content)
        return await _analyze_images_with_vlm(
            filename=filename,
            mime_type=mime_type,
            content=content,
            image_bytes=safe_bytes,
            image_mime=safe_mime,
            user_query=user_query,
        )

    heuristics = classify_extracted_text(text)
    return AttachmentAnalysisResult(
        filename=filename,
        mime_type=mime_type,
        file_sha256=calculate_sha256(content),
        extraction_method="text-layer",
        extracted_text=text[: settings.ATTACHMENT_CONTEXT_MAX_CHARS],
        structured_data={**heuristics, "detected_type": heuristics["detected_type"]},
        verification_status="not_applicable",
        confidence=1.0 if text.strip() else 0.0,
        warnings=[] if text.strip() else ["empty_text_layer"],
        processing_ms=int((time.time() - start_time) * 1000),
    )


def extract_text_from_file(file_bytes: bytes, filename: str, content_type: str | None = None) -> str:
    ext = validate_file_type(filename, content_type, file_bytes)
    if ext in TEXT_EXTENSIONS:
        return _extract_from_text(file_bytes)[: settings.ATTACHMENT_CONTEXT_MAX_CHARS]
    raise RuntimeError("Ekstraksi file non-teks berjalan melalui attachment processor.")


class AttachmentContextPackage(BaseModel):
    user_question: str = ""
    file_summary: str
    ocr_text: str
    detected_type: str
    extracted_entities: list[dict] = Field(default_factory=list)
    verified_candidates: list[dict] = Field(default_factory=list)
    neo4j_evidence: list[dict] = Field(default_factory=list)
    confidence: float = 0.0
    verification_status: str = "unverified"
    limitations: list[str] = Field(default_factory=list)


def build_attachment_context_package(analysis: AttachmentAnalysisResult, *, user_question: str = "") -> AttachmentContextPackage:
    detected_type = analysis.structured_data.get("detected_type", "unknown")
    limitations = list(analysis.structured_data.get("limitations") or [])
    limitations.extend(analysis.warnings or [])
    if analysis.verification_status not in {"verified", "not_applicable"}:
        limitations.append("Identitas molekul atau tanaman belum boleh dianggap pasti karena evidence verifikasi belum kuat.")
    if detected_type == "chemical_structure_diagram":
        limitations.append("VLM hanya membaca bukti visual; sistem belum melakukan Optical Chemical Structure Recognition tervalidasi menjadi SMILES/InChI.")
    deduped_limitations: list[str] = []
    for item in limitations:
        if item and item not in deduped_limitations:
            deduped_limitations.append(item)
    return AttachmentContextPackage(user_question=user_question, file_summary=f"{analysis.filename} diproses via {analysis.extraction_method}.", ocr_text=analysis.extracted_text, detected_type=detected_type, extracted_entities=analysis.detected_entities, verified_candidates=analysis.neo4j_candidates, neo4j_evidence=analysis.neo4j_candidates, confidence=analysis.confidence, verification_status=analysis.verification_status, limitations=deduped_limitations)


def format_attachment_context_package(package: AttachmentContextPackage) -> str:
    entities = "\n".join(f"- {entity.get('type')}: {entity.get('value')}" for entity in package.extracted_entities[:30]) or "Tidak ada entitas eksplisit yang terdeteksi."
    candidates = "\n".join(f"- {candidate.get('name', 'unknown')} | score={float(candidate.get('score') or 0):.2f} | evidence={', '.join(candidate.get('matched_evidence') or [])}" for candidate in package.verified_candidates[:8]) or "Tidak ada kandidat Neo4j yang cukup kuat."
    limitations = "\n".join(f"- {item}" for item in package.limitations[:12]) or "Tidak ada."
    return f"""[ATTACHMENT EVIDENCE]
Filename:
{package.file_summary}

Detected content type:
{package.detected_type}

Visual/Text evidence:
{package.ocr_text[:settings.ATTACHMENT_CONTEXT_MAX_CHARS]}

Entities extracted:
{entities}

Neo4j candidates:
{candidates}

Verification status:
{package.verification_status}

Confidence:
{package.confidence:.2f}

Limitations:
{limitations}
[/ATTACHMENT EVIDENCE]"""
